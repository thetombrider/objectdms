# Standard library imports
import io
from typing import Dict, Any

# Third-party imports
from PIL import Image
from PIL.ExifTags import TAGS
import PyPDF2
from docx import Document as DocxDocument
import magic
from fastapi import UploadFile

class MetadataExtractor:
    """
    Service for extracting metadata from various file types.
    
    This class provides methods to extract metadata from different file formats:
    - Images (JPEG, PNG, etc.): Extracts EXIF data, dimensions, format
    - PDFs: Extracts page count, author, creation date, etc.
    - Word documents: Extracts author, creation date, revision, etc.
    - Other files: Extracts basic metadata like size and MIME type
    """
    
    async def extract_metadata(self, file: UploadFile, mime_type: str) -> Dict[str, Any]:
        """
        Extract metadata based on file type.
        
        Args:
            file: The uploaded file to process
            mime_type: The MIME type of the file
            
        Returns:
            Dict[str, Any]: Extracted metadata
        """
        # Read file content
        content = await file.read()
        file.file.seek(0)  # Reset file position for subsequent reads
        
        # Route to appropriate extractor based on MIME type
        if mime_type.startswith('image/'):
            return self._extract_image_metadata(content)
        elif mime_type == 'application/pdf':
            return self._extract_pdf_metadata(content)
        elif mime_type in [
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]:
            return self._extract_docx_metadata(content)
        else:
            return self._extract_basic_metadata(content)

    def _extract_image_metadata(self, content: bytes) -> Dict[str, Any]:
        """
        Extract metadata from image files.
        
        Features:
        - Basic image properties (format, mode, dimensions)
        - EXIF data if available
        - Handles various image formats
        
        Args:
            content: Raw image file content
            
        Returns:
            Dict[str, Any]: Image metadata
        """
        try:
            image = Image.open(io.BytesIO(content))
            metadata = {
                "format": image.format,
                "mode": image.mode,
                "size": {"width": image.width, "height": image.height},
            }
            
            # Extract EXIF data if available
            if hasattr(image, '_getexif') and image._getexif():
                exif = image._getexif()
                for tag_id in exif:
                    tag = TAGS.get(tag_id, tag_id)
                    data = exif[tag_id]
                    # Convert bytes to string for JSON serialization
                    if isinstance(data, bytes):
                        data = data.decode(errors='ignore')
                    metadata[tag] = str(data)
            
            return metadata
        except Exception as e:
            return {"error": str(e)}

    def _extract_pdf_metadata(self, content: bytes) -> Dict[str, Any]:
        """
        Extract metadata from PDF files.
        
        Features:
        - Page count
        - Document properties (author, title, etc.)
        - Encryption status
        - Text preview from first page
        
        Args:
            content: Raw PDF file content
            
        Returns:
            Dict[str, Any]: PDF metadata
        """
        try:
            pdf = PyPDF2.PdfReader(io.BytesIO(content))
            info = pdf.metadata
            metadata = {
                "pages": len(pdf.pages),
                "encrypted": pdf.is_encrypted,
            }
            
            # Extract document info if available
            if info:
                for key, value in info.items():
                    if isinstance(value, bytes):
                        value = value.decode(errors='ignore')
                    metadata[key] = str(value)
            
            # Extract preview text from first page
            if len(pdf.pages) > 0:
                first_page = pdf.pages[0]
                metadata["preview"] = first_page.extract_text()[:500]  # First 500 chars
            
            return metadata
        except Exception as e:
            return {"error": str(e)}

    def _extract_docx_metadata(self, content: bytes) -> Dict[str, Any]:
        """
        Extract metadata from Word documents.
        
        Features:
        - Document structure (sections, paragraphs)
        - Core properties (author, created, modified)
        - Text preview from first paragraphs
        
        Args:
            content: Raw Word document content
            
        Returns:
            Dict[str, Any]: Document metadata
        """
        try:
            doc = DocxDocument(io.BytesIO(content))
            metadata = {
                "sections": len(doc.sections),
                "paragraphs": len(doc.paragraphs),
            }
            
            # Extract core properties
            core_properties = doc.core_properties
            if core_properties:
                metadata.update({
                    "author": str(core_properties.author),
                    "created": str(core_properties.created),
                    "modified": str(core_properties.modified),
                    "title": str(core_properties.title),
                    "revision": str(core_properties.revision),
                })
            
            # Generate text preview
            preview_text = ""
            for para in doc.paragraphs[:5]:  # First 5 paragraphs
                preview_text += para.text + "\n"
            metadata["preview"] = preview_text[:500]  # First 500 chars
            
            return metadata
        except Exception as e:
            return {"error": str(e)}

    def _extract_basic_metadata(self, content: bytes) -> Dict[str, Any]:
        """
        Extract basic metadata for unsupported file types.
        
        Features:
        - File size
        - Detected MIME type
        
        Args:
            content: Raw file content
            
        Returns:
            Dict[str, Any]: Basic metadata
        """
        try:
            mime = magic.from_buffer(content[:2048], mime=True)
            return {
                "detected_mime_type": mime,
                "size": len(content),
            }
        except Exception as e:
            return {"error": str(e)}

# Global instance
metadata_extractor = MetadataExtractor() 